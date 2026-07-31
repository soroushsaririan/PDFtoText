from pathlib import Path

from docx import Document

from pdf_form_kit import FormPage
from sample_data_content import PARTICIPANTS

ROOT = Path(__file__).resolve().parent.parent
OUT_DIR = ROOT / "sample_usb"

CUE_OPTIONS = [
    "Faster cadence", "Soft quiet steps", "Slight forward trunk lean",
    "Squeezing the gluteal muscles", "Tucking in the chin",
    "Keeping arms linear", "Abdominal tension", "Landing feet under hips",
    "Taking breaks when needed",
]
EXERCISE_OPTIONS = [
    "Hip Strength", "Foot Strength", "Core Strength", "Plyometric",
    "Neuromuscular Control", "Stretching", "Other",
]

HAND_FONT_BY_PID = {
    "VPT001": "caveat", "VPT002": "shadows", "VPT003": "sacramento",
    "VPT004": "sacramento", "VPT005": "caveat",
}


def write_field(page, label, key, wk, messiness):
    correction = wk.get("corrections", {}).get(key)
    if correction:
        wrong, correct = correction
        page.scratched_field(label, wrong, correct, messiness=messiness)
    else:
        page.field(label, wk[key], messiness=messiness)


def pain_fields(page, wk, messiness=1.0):
    page.section("Musculoskeletal Pain")
    write_field(page, "Pain RIGHT NOW (0-10 scale):", "pain_now", wk, messiness)
    page.field("USUAL pain during the last week:", wk["pain_usual"], messiness=messiness)
    page.field("BEST pain during the last week:", wk["pain_best"], messiness=messiness)
    page.field("WORST pain during the last week:", wk["pain_worst"], messiness=messiness)


def build_intake_crf(path, pid, profile, hand_font, seed, messiness=1.0):
    page = FormPage(path, hand_font=hand_font, seed=seed)
    page.title("Virtual Physical Therapy Support for Runners")
    page.subtitle("Visit 1 Case Report Form")
    page.field("Participant ID:", pid, messiness=messiness)
    page.field("Tester:", profile["tester"], messiness=messiness)
    page.section("Characteristics")
    page.yes_no("Sex (Yes = Female):", "Yes" if profile["sex"] == "Female" else "No")
    page.field("Age (yrs):", profile["age"], messiness=messiness)
    page.section("Musculoskeletal Pain")
    page.field("Injury diagnosis (if known):", profile["injury"], messiness=messiness)
    page.field("Current injury location:", profile["injury"].split(",")[-1].strip(), messiness=messiness)
    page.field("Have you seen a doctor about this problem?", "Yes", messiness=messiness)
    page.save()


def build_weekly_crf(path, pid, wk, hand_font, seed, messiness=1.0):
    page = FormPage(path, hand_font=hand_font, seed=seed)
    page.title("Weekly Virtual Physical Therapy (VPT) Session")
    page.subtitle("Case Report Form")
    page.field("Participant ID:", pid, messiness=messiness)
    page.field("Date:", wk["date"], messiness=messiness)
    page.field("Weekly visit #:", str(wk["week"]), messiness=messiness)
    page.section("Adherence to Home Exercise Program")
    page.field("Days completed exercise as prescribed:", wk["days_exercised"], messiness=messiness)
    page.field("Number of exercises per session:", wk["exercises_per_session"], messiness=messiness)
    page.checkboxes("Exercise types performed:", EXERCISE_OPTIONS, wk["exercise_types"], messiness=messiness)
    page.yes_no("Any exercises too difficult?", wk["too_difficult"])
    page.yes_no("Confident in correct running form?", wk["confident_form"])
    page.yes_no("Incorporated running-form cues?", wk["cues_used"])
    if wk["cues_used"] == "Yes" and wk["cues_list"]:
        page.checkboxes("Which cues:", CUE_OPTIONS, wk["cues_list"], messiness=messiness)
    page.section("Participation in Running & Other Exercise")
    page.field("Days ran this week:", wk["days_ran"], messiness=messiness)
    page.field("Mileage this week:", wk["mileage"], messiness=messiness)
    write_field(page, "Total minutes running this week (min):", "running_minutes", wk, messiness)
    page.yes_no("Other exercise this week?", wk["other_exercise"])
    if wk["other_exercise_desc"]:
        page.note("Describe:", wk["other_exercise_desc"], messiness=messiness)
    pain_fields(page, wk, messiness=messiness)
    page.note("Therapist / participant notes:", wk["notes"], lines=3, messiness=messiness)
    page.save()


def build_final_crf(path, pid, final, hand_font, seed, messiness=1.0):
    page = FormPage(path, hand_font=hand_font, seed=seed)
    page.title("Virtual Physical Therapy Support for Runners")
    page.subtitle("Final Visit (week 25) Case Report Form")
    page.field("Participant ID:", pid, messiness=messiness)
    page.section("Adherence to Home Exercise Program")
    page.field("Days completed exercise (days/week):", final["days_exercised"], messiness=messiness)
    page.field("Number of exercises per session:", final["exercises_per_session"], messiness=messiness)
    page.yes_no("Confident in correct running form?", final["confident_form"])
    page.field("Mileage this week:", final["mileage"], messiness=messiness)
    page.section("Musculoskeletal Pain")
    page.field("Pain RIGHT NOW (0-10 scale):", final["pain_now"], messiness=messiness)
    page.field("USUAL pain during the last week:", final["pain_usual"], messiness=messiness)
    page.field("BEST pain during the last week:", final["pain_best"], messiness=messiness)
    page.field("WORST pain during the last week:", final["pain_worst"], messiness=messiness)
    page.note("Overall summary / therapist notes:", final["notes"], lines=3, messiness=messiness)
    page.save()


def build_exit_survey(path, pid, exit_data, hand_font, seed, messiness=1.0):
    page = FormPage(path, hand_font=hand_font, seed=seed)
    page.title("Virtual Physical Therapy - Exit Survey")
    page.field("Participant ID:", pid, messiness=messiness)
    page.field("Overall satisfaction with the program:", exit_data["satisfaction"], messiness=messiness)
    page.yes_no("Would you recommend this program?", exit_data["would_recommend"])
    page.note("Additional comments:", exit_data["comments"], lines=3, messiness=messiness)
    page.save()


def build_missed_week_docx(path, pid, week, note):
    doc = Document()
    doc.add_paragraph(f"{pid} No Week {week} VPT Session")
    doc.add_paragraph(note)
    doc.save(str(path))


def placeholder_dir(path, note):
    path.mkdir(parents=True, exist_ok=True)
    (path / "README.txt").write_text(
        note + "\n\nThis sample project only analyzes handwritten PDF forms; "
        "gait-report / video folders like this one are ignored by the pipeline.\n"
    )


def build_participant(pid, data, seed_base):
    profile = data["profile"]
    hand_font = HAND_FONT_BY_PID.get(pid, "caveat")
    messiness = data.get("messiness", 1.0)
    pdir = OUT_DIR / pid
    pdir.mkdir(parents=True, exist_ok=True)

    build_intake_crf(pdir / f"{pid} CRF visit 1.pdf", pid, profile, hand_font, seed_base + 1, messiness=messiness)
    build_final_crf(pdir / f"{pid} CRF visit 3.pdf", pid, data["final"], hand_font, seed_base + 3, messiness=messiness)
    build_exit_survey(pdir / f"{pid} Exit survey.pdf", pid, data["exit"], hand_font, seed_base + 4, messiness=messiness)

    log_dir = pdir / f"{pid} weekly exercise log"
    for i, wk in enumerate(data["weeks"]):
        week_dir = log_dir / f"Week {wk['week']}"
        week_dir.mkdir(parents=True, exist_ok=True)
        if wk.get("missed"):
            build_missed_week_docx(
                week_dir / f"{pid} No Week{wk['week']} VPT session.docx",
                pid, wk["week"], wk["notes"],
            )
        else:
            build_weekly_crf(
                week_dir / f"{pid} Week{wk['week']} VPT session.pdf",
                pid, wk, hand_font, seed_base + 10 + i, messiness=messiness,
            )

    for visit in (1, 2, 3):
        placeholder_dir(pdir / f"DARI FUNCTION visit {visit}", f"{pid} gait function data placeholder, visit {visit}.")
        placeholder_dir(pdir / f"DARI RUNNING visit {visit}", f"{pid} treadmill running data placeholder, visit {visit}.")
    placeholder_dir(pdir / f"{pid} Video visit 1", f"{pid} session video placeholder.")


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    for i, (pid, data) in enumerate(PARTICIPANTS.items()):
        build_participant(pid, data, seed_base=100 * (i + 1))
        print(f"Generated sample data for {pid}")
    print(f"\nSample USB folder ready at: {OUT_DIR}")


if __name__ == "__main__":
    main()
